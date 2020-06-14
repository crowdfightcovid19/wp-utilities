# -*- coding: utf-8 -*-
"""
To use this code you need to have credentials for google sheets and google
drive api's. You need
to create an environmental variable called CRF_CREDENTIALS pointing to the
folder containing the json file with the credentials. To create that variable, execute

os.environ['CRF_CREDENTIALS'] = r'folder_that_contains_your_credentials'

You must also have an environmental variable pointing to the folder where you 
store your crowdfight-related data. To create this variable, execute

os.environ['CRF_DATA'] = r'folder_that_contains_your_data'

This program also needs the template xxx_report.docx, which must be in the same
folder as the code

Created on Thu Jun 11 11:31:26 2020

@author: Alfonso
"""


import gspread
from oauth2client.service_account import ServiceAccountCredentials
from apiclient import discovery
from shutil import copyfile
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import win32api
import win32print
import time
import os

def test():
    """
    The links are generated by selecting one or several files in gdrive,
    right-click, share, copy link. Then just paste the result.
    """
    generate_reports('https://docs.google.com/spreadsheets/d/1IQQ9toViwCb4dEF50Zlc7wWJNsh1ioZTi5lOmxBzyMs/edit?usp=sharing,%20https://docs.google.com/spreadsheets/d/1Rje2i6swuzbURXBZIy7x_QbAqwo-BnEI9e0-dtNtwKg/edit?usp=sharing')

def cosa():
    print('—')
    print(chr(8212))

def generate_reports(url_batch,service=[],sheet_api=[],drive_api=[],db=[]):    
    if not(not(service)):
        sheet_api = service['sheet_api']
        drive_api = service['drive_api']
        db = service['db']
    if not(sheet_api) or not(drive_api):        
        sheet_api, drive_api = open_apis()
    if not(db):        
        db = load_db(sheet_api)
    url_list = url_batch.split(',%20')
    for url in url_list:
        info_request = sheet_to_inforequest(sheet_api,url)
        print(info_request['request_number'])        
        complete_inforequest(db, info_request)        
        write_report(info_request)     
        update_organizertable(sheet_api,info_request)           
    service = {
        'sheet_api' : sheet_api,
        'drive_api' : drive_api,
        'db' : db}
    return service

def update_organizertable(sheet_api,info_request): # Update organization table
    sheet_organizer = sheet_api.open_by_url(\
        'https://docs.google.com/spreadsheets/d/1JCtFrOjdDdxI99NPGoMfWO9TWiVGOmVHE9Dqg8Yxi8c/edit#gid=906247198'\
         ).worksheet('Ready to publish')
    request_number_list = sheet_organizer.col_values(1)[1:]
    try:
        ind_row = request_number_list.index(info_request['request_number'][0]) + 2
    except:
        ind_row = len(request_number_list) + 2 # If the request is not yet in the table, add it at the end
    # print([info_request['request_number'][0],info_request['url_table'][0],
    #       datetime.now().strftime('%Y%m%dT%H%M%S')])
    sheet_organizer.update(
        'A'+str(ind_row)+':C'+str(ind_row),\
         [[info_request['request_number'][0],info_request['url_table'][0],
          datetime.now().strftime('%Y%m%dT%H%M%S')]]
         )

def open_apis():
    # use creds to create a client to interact with the Google Drive API
    #scope = ['https://spreadsheets.google.com/feeds']
    scope = ['https://spreadsheets.google.com/feeds',
             'https://www.googleapis.com/auth/drive']
    creds = ServiceAccountCredentials.from_json_keyfile_name(os.path.join(os.environ['CRF_CREDENTIALS'],'google.json'), scope)
    sheet_api = gspread.authorize(creds)
    drive_api = discovery.build('drive', 'v3', credentials=creds)
    return sheet_api, drive_api

def load_db(sheet_api): # Load the common databases (volunteer database, staff database, etc.)
    db = {}
    # Load staff database
    book_advisormanagement = sheet_api.open_by_url('https://docs.google.com/spreadsheets/d/1tKaNv6P1vwv1ll21GBmeKNYSNI11N_kTZiW2yjfqAW0/edit#gid=1858226365')
    sheet_staff = book_advisormanagement.get_worksheet(-1)
    db['staff'] = sheet_staff.get_all_values()
    # Load volunteer database
    sheet_volunteer = sheet_api.open_by_url('https://docs.google.com/spreadsheets/d/15kBVUpcsb1Y3vAaHYFV4WnTjNUO6wiBXhuciJUEAx-I/edit#gid=1550850176').sheet1
    db['volunteer'] = sheet_volunteer.get_all_values()
    return db
    

def sheet_to_inforequest(sheet_api,url_sheet): # Reads information from the spreadsheet    
    sheet_origin = sheet_api.open_by_url(url_sheet).sheet1
    list_origin = sheet_origin.get_all_values()
    info_request = {}
    info_request['url_table'] = [url_sheet]
    info_nested = {}
    name_nest = ''
    for row in list_origin:
        if ''.join(row)=='': # If empty row, we are moving from one section to the next
            if name_nest!='': # Save previous nested dict (if it exists), and reset nested dict
                info_request[name_nest] = info_nested
                name_nest = ''
                info_nested = {}       
        if row[0]!='': # New section
            if name_nest!='': # Save previous nested dict (if it exists) - This should not be necessary if all sections are separated by blank lines
                info_request[name_nest] = info_nested
            # Reset nested dict
            info_nested = {}
            name_nest = row[0].replace(' ','_').replace('?','').lower()
            if name_nest=='volunteers':
                name_nest = 'volunteer' # For consistency with coding style
            if name_nest=='collaborators':
                name_nest = 'collaborator' # For consistency with coding style
        if row[1]!='': # There is content in this line
            ind_nonempty_last = [i for i,e in enumerate(row) if e][-1]   
            if row[1].replace(' ','_').replace('?','').lower()=='email': # If we are handling e-mails, remove spaces
                row[2:ind_nonempty_last+1] = [item.replace(' ','') for item in row[2:ind_nonempty_last+1]]
            if name_nest=='': # If not nested, just add it to the main dict
                info_request[row[1].replace(' ','_').replace('?','').lower()] = row[2:ind_nonempty_last+1]
            else: # If nested, add to the nested dict
                info_nested[row[1].replace(' ','_').replace('?','').lower()] = row[2:ind_nonempty_last+1]
    if name_nest!='': # If there is a nested dict unsaved, save it
        info_request[name_nest] = info_nested
    # Now do the dates and times
    column = [row[1] for row in list_origin]
    tag_date_list = list_origin[column.index('Dates and times:')]
    date_list = list_origin[column.index('Dates and times:')+1]
    dict_date = {}
    for i_date in [2,3,4,5,6]:
        dict_date[tag_date_list[i_date].replace(' ','_').replace('1','fir').lower()] = date_list[i_date:i_date+1] # I leave it inside a list for consistency with other fields
    info_request['date']= dict_date
    
    # Fill empty fields with ''
    for key in info_request.__iter__():
        if isinstance(info_request[key],dict):
            for key2 in info_request[key].__iter__():
                if not(info_request[key][key2]):
                    info_request[key][key2] = ['']                
        else:
            if not(info_request[key]):
                info_request[key] = [''] 
    
    return info_request

def complete_inforequest(db,info_request): # Completes the information in info_request using our databases        
    # Complete information about the staff
    email_staff_list = [row[1].replace(' ','') for row in db['staff']]
    type_staff_list = ['receiver', 'advisor_manager', 'advisor', 'coordinator', 'documenter', 'follow_up']
    field_staff_list = ['full_name', '', '', 'institution', 'city', 'country', 'full_affiliation', 'wants_to_be_named_publicly'] # indicates the position of each type of info in the database  
    for type_staff in type_staff_list:
        for i_person,email in enumerate(info_request[type_staff]['email']):
            if email.lower()!='none' and not(not(email)):
                ind_staff = email_staff_list.index(email)
                for field_info in info_request[type_staff].keys():
                    if info_request[type_staff][field_info]!='email':
                        if len(info_request[type_staff][field_info])<i_person+1:
                            info_request[type_staff][field_info].append('')
                        if not(info_request[type_staff][field_info][i_person]):
                            info_request[type_staff][field_info][i_person] = db['staff'][ind_staff][field_staff_list.index(field_info)]    
        
    # Load volunteer database
    email_list = [row[1].replace(' ','') for row in db['volunteer']]
    type_person_list = ['requester','volunteer','collaborator']
    field_info_list = ['full_name', 'institution'] # indicates the position of each type of info in the database
    ind_info_db_list = [2,3]
    for type_person in type_person_list:
        for i_person,email in enumerate(info_request[type_person]['email']):
            if email.lower()!='none' and not(not(email)):
                try:
                    ind_person_db = email_list.index(email)
                    #print(db['volunteer'][ind_person_db])
                except:
                    ind_person_db = -1
                    if type_person=='volunteer':
                        print(email + ' not found in volunteer database')                
                for [field_info,ind_info_db] in zip(field_info_list,ind_info_db_list):                    
                    if len(info_request[type_person][field_info])<i_person+1:
                        info_request[type_person][field_info].append('')
                    if not(info_request[type_person][field_info][i_person]) and ind_person_db>=0:
                        info_request[type_person][field_info][i_person] = \
                            db['volunteer'][ind_person_db][ind_info_db]   
    
    # Remove fields flagged for removal
    for key in info_request.__iter__():
        if isinstance(info_request[key],dict):
            for key2 in info_request[key].__iter__():
                for i in range(len(info_request[key][key2])):
                    if info_request[key][key2][i].lower()=='<remove>':  
                        #print(key)
                        #print(key2)
                        #print(info_request[key]['full_name'][i])
                        info_request[key][key2][i]=''
        else:
            for i in range(len(info_request[key])):
                if info_request[key][i].lower()=='<remove>':                    
                    info_request[key][i]=''
    
    # Complete full affiliations
    type_person_list = ['receiver', 'advisor_manager', 'advisor', 'coordinator', 'documenter', 'follow_up','requester','volunteer','collaborator']                         
    for type_person in type_person_list:
        #print(type_person)
        for i_person in range(len(info_request[type_person]['email'])):
            if info_request[type_person]['email'][i_person]!='none':
                if len(info_request[type_person]['full_affiliation'])<i_person+1:
                    info_request[type_person]['full_affiliation'].append('')
                if info_request[type_person]['full_affiliation'][i_person]=='':
                    if len(info_request[type_person]['institution'])>i_person \
                        and info_request[type_person]['institution'][i_person]!='':
                        info_request[type_person]['full_affiliation'][i_person] = \
                            info_request[type_person]['institution'][i_person] + '<'
                    if len(info_request[type_person]['city'])>i_person \
                        and info_request[type_person]['city'][i_person]!='':    
                        info_request[type_person]['full_affiliation'][i_person] = \
                            info_request[type_person]['full_affiliation'][i_person] \
                            + '>' + info_request[type_person]['city'][i_person] + '<'                        
                    if len(info_request[type_person]['country'])>i_person \
                        and info_request[type_person]['country'][i_person]!='':    
                        info_request[type_person]['full_affiliation'][i_person] = \
                            info_request[type_person]['full_affiliation'][i_person] \
                            + '>' + info_request[type_person]['country'][i_person]      
                    info_request[type_person]['full_affiliation'][i_person] = \
                        info_request[type_person]['full_affiliation'][i_person].replace('<>',', ')
                    info_request[type_person]['full_affiliation'][i_person] = \
                        info_request[type_person]['full_affiliation'][i_person].replace('<','')
                    info_request[type_person]['full_affiliation'][i_person] = \
                        info_request[type_person]['full_affiliation'][i_person].replace('>','')

def write_name_and_affiliation(info_request_typeperson,ind_person):    
    #print(info_request_typeperson['wants_to_be_named_publicly'][ind_person].lower())
    #print(info_request_typeperson)
    if info_request_typeperson['email'][ind_person]!='none':        
        if len(info_request_typeperson['wants_to_be_named_publicly'])>ind_person \
            and info_request_typeperson['wants_to_be_named_publicly'][ind_person].lower()=='yes':
            text_person = info_request_typeperson['full_name'][ind_person]                
        else:
            text_person = 'Anonymous'
        if not(not(info_request_typeperson['full_affiliation'][ind_person])):
            text_person = text_person + ' | ' + info_request_typeperson['full_affiliation'][ind_person]
    else:
        text_person = 'None'
    return text_person
    
# Writes report in Word
def write_report(info_request,print_to_pdf=True,\
                 path_report_folder=os.path.join(os.environ['CRF_DATA'],'Request_reports')):     
    if info_request['table_finished'][0].lower()!='yes':
        print('WARNING!! Table is marked as NOT FINISHED.')
    if info_request['final_validation_done'][0].lower()!='yes':
        print('WARNING!! Table is marked as NOT VALIDATED.')
    footnote_list = []
    now = datetime.now()
    name_file_report = info_request['request_number'][0] + '_Report_' + now.strftime('%Y%m%dT%H%M%S')
    name_file_emails = info_request['request_number'][0] + '_Report_' + now.strftime('%Y%m%dT%H%M%S') + '_emails'
    path_file_report = os.path.join(path_report_folder,name_file_report + '.docx')
    path_file_emails = os.path.join(path_report_folder,name_file_emails + '.txt')
    copyfile(r'xxx_report.docx',path_file_report)
    document = Document(path_file_report)
    document.styles['Normal'].font.name = 'Raleway'
    document.styles['Normal'].font.size = Pt(12)
    
    if info_request['final_outcome'][0]=='Resolved: Task successfully finished':
        info_request['final_outcome'][0] = 'Success'
    
    # Title
    p = document.add_paragraph()
    run = p.add_run('Request ' + info_request['request_number'][0])
    run.font.bold = True
    run.font.size = Pt(16)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Basic info about the request
    p = document.add_paragraph()
    p = document.add_paragraph()
    p.add_run('Request description: ' ).bold = True
    p.add_run(info_request['description_task'][0])
    p = document.add_paragraph()
    p.add_run('Requester\'s institution: ' ).bold = True
    p.add_run(info_request['requester']['full_affiliation'][0])
    p = document.add_paragraph()
    p.add_run('Outcome: ' ).bold = True
    p.add_run(info_request['final_outcome'][0] + chr(8212) + info_request['description_success'][0])    
    p = document.add_paragraph()
    
    # Volunteers    
    p = document.add_paragraph()
    p.add_run('Volunteers:').bold = True
    footnote_list_role = []
    # any_volunteer = False
    any_role = False
    for i_volunteer in range(len(info_request['volunteer']['email'])):
        # if info_request['volunteer']['email'][i_volunteer]!='none':
        #     any_volunteer = True
        p = document.add_paragraph()
        if len(info_request['volunteer']['role'])>i_volunteer \
           and not(not(info_request['volunteer']['role'][i_volunteer])):
            add_footnote(p,footnote_list_role,info_request['volunteer']['role'][i_volunteer],type_symbol='symbol')
            any_role = True
        p.add_run(write_name_and_affiliation(info_request['volunteer'], i_volunteer))
    # Volunteers' roles
    if any_role:
        p = document.add_paragraph()
        p = document.add_paragraph()
        p.add_run('Volunteers\' roles:').bold = True
        print_footnotes(document,footnote_list_role,font_size=12)
        
    
    # Timeline
    p = document.add_paragraph()
    p = document.add_paragraph()
    p.add_run('Timeline:').bold = True    
    if info_request['date']['request'][0]!='none':
        p = document.add_paragraph('Date of request: ' \
                                   +  datetime.strptime(info_request['date']['request'][0],'%d/%m/%Y %H:%M').strftime('%B %d, %Y'))
        deltaT = datetime.strptime(info_request['date']['first_response'][0],'%d/%m/%Y %H:%M') \
            - datetime.strptime(info_request['date']['request'][0],'%d/%m/%Y %H:%M')
        p = document.add_paragraph(deltaT_to_string(deltaT) + ' from request to first response')
        add_footnote(p,footnote_list,'The first response from Crowdfight typically contains a feasibility estimate for the request.')
        deltaT = datetime.strptime(info_request['date']['proposal'][0],'%d/%m/%Y %H:%M') \
            - datetime.strptime(info_request['date']['first_response'][0],'%d/%m/%Y %H:%M')
        p = document.add_paragraph(deltaT_to_string(deltaT) + ' from first response to proposal')
        add_footnote(p,footnote_list,'The proposal typically includes making first contact between requester and volunteers.')
        deltaT = datetime.strptime(info_request['date']['success'][0],'%d/%m/%Y %H:%M') \
            - datetime.strptime(info_request['date']['proposal'][0],'%d/%m/%Y %H:%M')
        p = document.add_paragraph(deltaT_to_string(deltaT) + ' from proposal to resolution')
    else:
        p = document.add_paragraph('Date of request: None')                
    p = document.add_paragraph('Date of resolution') 
    add_footnote(p,footnote_list,'For successful requests, this is the date in which we confirmed that the success criterion was reached (we try to be unintrusive during follow-up, so this confirmation can have significant delay). The success criterion is defined beforehand, and is as stringent as possible to guarantee that our activity is truly useful.')
    add_footnote(p,footnote_list,'For successful requests, this is the date in which we confirmed that the success criterion was reached (we try to be unintrusive during follow-up, so this confirmation can have significant delay). The success criterion is defined beforehand, and is as stringent as possible to guarantee that our activity is truly useful.')
    p.add_run(': ' +  datetime.strptime(info_request['date']['success'][0],'%d/%m/%Y %H:%M').strftime('%B %d, %Y'))
    
    # Crowdfight staff
    p = document.add_paragraph()    
    type_person_list = [['receiver','Receiver'],
                        ['advisor_manager','Advisor manager'],
                        ['advisor','Scientific advisor'],
                        ['coordinator','Coordinator'],                        
                        ['follow_up','Follow up'],
                        ['documenter','Documenter']
                        ]
    text_footnote_person = ['The Receiver is in charge of receiving the request, understanding it (with help of other coordinators and advisors), and writing the request to be sent in the task distribution e-mail (when appropriate).',
                            'The advisor manager handles the assignment of requests to the scientific advisors',
                            'The scientific advisor reviews the self-assessments submitted by the volunteers, shortlists the most adequate ones for the request, and recommends a course of action to the coordinator.',
                            'The coordinator decides a course of action with help from the scientific advisor, makes a proposal to the requester and contacts the selected volunteers.',
                            'The person in charge of follow-up stays in touch with requester and volunteers to make sure they have everything they need until the resolution of the request.',
                            'The documenter maintains our database and gathers the necessary information to produce this report.'
                            ]
    
    for i_person,person in enumerate(type_person_list):
        p = document.add_paragraph()
        p.add_run(person[1]).bold = True
        add_footnote(p,footnote_list,text_footnote_person[i_person],bold=True)
        p.add_run(': ').bold = True
        # print(person)
        # print(write_name_and_affiliation(info_request[person[0]], 0))
        p.add_run(write_name_and_affiliation(info_request[person[0]], 0))
        # if info_request[person[0]]['email'][0]!='none':
        #     p.add_run(info_request[person[0]]['full_name'][0] + ' | ' + info_request[person[0]]['full_affiliation'][0])
        # else:
        #     p.add_run('None')
    
    # Notes
    p = document.add_paragraph()    
    p = document.add_paragraph()    
    p.add_run('Notes:').bold = True
    print_footnotes(document,footnote_list)
    p = document.add_paragraph()    
    p = document.add_paragraph()    
    if not(info_request['doi']['static']):
        doi_static = 'xxxx'
    else:
        doi_static = info_request['doi']['static'][-1] # Take the last (we'll keep adding as we make new versions)
    if not(info_request['doi']['updated']):
        doi_updated = 'xxxx'
    else:
        doi_updated = info_request['doi']['updated'][0]
    run = p.add_run('This report belongs to the internal documentation of Crowdfight, a non-for profit association registered in the EU with number ES-G47805569. Reports are published for transparency with the explicit agreement of all people mentioned. The report was generated automatically from our records on 08-Jun-2020 15:23:09. While we do our best to keep faithful records, it may contain unintentional inaccuracies or omissions. Any modification of this document is strictly forbidden. This document can be found at https://doi.org/' \
                    + doi_static + ', and up to date versions can be found at https://doi.org/' \
                    + doi_updated + '.')
    run.font.size = Pt(9)
    document.save(path_file_report)
    print_emails(info_request,file_path=path_file_emails)
    if print_to_pdf:
        printout_file(path_file_report,wait_for_pdf=True)

def add_footnote(p,footnote_list,text,bold=False,type_symbol='number'):
    symbol_list = ['*',chr(8224),chr(8225),chr(167),'#','**',
                   chr(8224)+chr(8224),chr(8225)+chr(8225),chr(167)+chr(167),'##']
    if not(footnote_list):
        footnote_list.append([])
        footnote_list.append([]) # First list for symbols, second for texts
    if footnote_list[1].count(text)==0:
        footnote_list[1].append(text)
        if type_symbol=='number':
            footnote_list[0].append(str(len(footnote_list[1])))
        else:
            footnote_list[0].append(symbol_list[len(footnote_list[1])-1])
    ind_footnote = footnote_list[1].index(text)
    super_text = p.add_run(footnote_list[0][ind_footnote])
    if footnote_list[0][ind_footnote]!='*':
        super_text.font.superscript = True      
    super_text.font.bold = bold

def print_footnotes(document,footnote_list,font_size=9):     
    footnote_list
    for i_footnote,text in enumerate(footnote_list[1]):
        p = document.add_paragraph()        
        super_text = p.add_run(footnote_list[0][i_footnote])
        if footnote_list[0][i_footnote]!='*': # Because the asterisk is already a superscript
            super_text.font.superscript = True    
        super_text.font.size = Pt(font_size)
        run = p.add_run(' ' + text)
        run.font.size = Pt(font_size)

def deltaT_to_string(deltaT):
    number = deltaT.days
    unit = 'days'
    if number==0:
        number = max(1,round(deltaT.seconds/3600))
        unit = 'hours'
    if number==1:
        unit = unit[:-1]
    string = str(number) + ' ' + unit
    return string

def print_emails(info_request,add_name=False,file_path=[]): # Print out the list of e-mails of the people who should be informed about the request
    type_person_list = [['receiver', 'advisor_manager', 'advisor', 'coordinator', 'documenter', 'follow_up'],['requester'],['volunteer'],['collaborator']]
    type_institution_list = ['Crowdfight','requester','volunteer','collaborator']
    if not(not(file_path)):
        text_file = open(file_path,'w')
    for i_institution,type_institution in enumerate(type_institution_list):
        if not(file_path):
            print(type_institution)
        else:
            text_file.write(type_institution + '\n')
        contact_list = []
        for type_person in type_person_list[i_institution]:
            for email,full_name in zip(info_request[type_person]['email'],info_request[type_person]['full_name']):
                if add_name:
                    contact_list.append(full_name + ' <' + email + '>')
                else:
                    contact_list.append(email)
        contact_list = set(contact_list)
        if not(file_path):
            print(', '.join(contact_list))
        else:
            text_file.write(', '.join(contact_list) + '\n')
        text_file.write('\n')
    
def printout_file(path_file,wait_for_pdf=False):    
    win32api.ShellExecute (
      0,
      "print",
      path_file,
      #
      # If this is None, the default printer will
      # be used anyway.
      #
      '/d:"%s"' % win32print.GetDefaultPrinter (),
      ".",
      0
    )
    # Wait until the pdf is created (and its size stops growing)
    filesize_last = 0
    while wait_for_pdf \
          and (
              not(os.path.isfile(path_file.replace('.docx','.pdf'))) 
               or filesize_last<os.path.getsize(path_file.replace('.docx','.pdf'))
               ):
        if not(not(os.path.isfile(path_file.replace('.docx','.pdf')))):
            filesize_last = os.path.getsize(path_file.replace('.docx','.pdf'))
        time.sleep(1)    
        